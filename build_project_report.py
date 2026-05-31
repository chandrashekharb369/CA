from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
REPORT_PATH = ROOT / "reports" / "CA_Intelligence_Suite_Project_Report.docx"

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
TEXT = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"
BORDER = "D7DBE2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def style_run(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", style=None, before=0, after=6, line=1.10, align=None, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, italic=italic, color=color or TEXT)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    style_run(r, bold=True, color=BLUE if level < 3 else NAVY)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        style_run(r, color=TEXT)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        style_run(r, color=TEXT)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = widths or [6.5 / len(headers)] * len(headers)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, HEADER_FILL)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(header))
        style_run(run, size=9, bold=True, color=TEXT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.text = ""
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            run = p.add_run(str(value))
            style_run(run, size=9, color=TEXT)
    para(doc, "", after=4)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.268)
    section.page_height = Inches(11.693)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0.875)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 14, BLUE, 20, 13),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    hdr = section.header.paragraphs[0]
    hdr.text = "CA Intelligence Suite Project Report"
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hdr.runs:
        style_run(run, size=8.5, color=MUTED)
    ftr = section.footer.paragraphs[0]
    ftr.text = "Prepared from project source code and artifacts"
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ftr.runs:
        style_run(run, size=8, color=MUTED)


def load_metrics():
    path = ROOT / "model_artifacts" / "metrics.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def dataset_overview():
    files = [
        ROOT / "train_dataset" / "financial_dataset_clean.csv",
        ROOT / "test_dataset" / "dataset_generic.csv",
        ROOT / "test_dataset" / "dataset_manufacturing.csv",
        ROOT / "test_dataset" / "dataset_service.csv",
        ROOT / "test_dataset" / "dataset_startup.csv",
        ROOT / "test_dataset" / "dataset_trading.csv",
    ]
    rows = []
    for file in files:
        if file.exists():
            df = pd.read_csv(file)
            counts = df["Category"].value_counts().to_dict() if "Category" in df.columns else {}
            rows.append([
                str(file.relative_to(ROOT)),
                f"{len(df):,}",
                f"{len(df.columns):,}",
                ", ".join(f"{k}: {v:,}" for k, v in counts.items()),
            ])
    return rows


PROJECT_FILES = [
    ("app.py", "Streamlit UI orchestrator. Handles theme, upload, tabs, predictions, dashboard, insights, and report download."),
    ("config.py", "Central constants for paths, labels, thresholds, training parameters, visualization palette, and app metadata."),
    ("generate_dataset.py", "Synthetic Schedule III transaction generator with GST, payment modes, category distributions, and controlled noise."),
    ("preprocess.py", "Training-time preprocessing: text normalization, TF-IDF, amount scaling, payment-mode encoding, and feature serialization."),
    ("train_model.py", "Dense TensorFlow/Keras classifier trained with backpropagation and an Isolation Forest anomaly model."),
    ("rule_engine.py", "Backward chaining engine that computes Profit and Loss, Balance Sheet, GST, reconciliation, and CA suggestions."),
    ("modules/ml_model.py", "Runtime artifact loading and hybrid ML plus keyword-rule fallback classification."),
    ("modules/financial_engine.py", "Financial summary, health classification, cash-flow proxy, and comparative analysis."),
    ("modules/compliance_engine.py", "Compliance score, tax provision validation, balance-sheet validation, and insight generation."),
    ("modules/anomaly_detection.py", "Statistical and ML anomaly detection with summary counts and review signals."),
    ("modules/visualization.py", "Plotly charts for Sankey flow, network graph, sunburst hierarchy, trend, GST, and compliance gauge."),
    ("modules/report_generator.py", "ReportLab-based 14-section CA PDF report generator used by the Streamlit Report tab."),
    ("rules/compliance_rules.json", "JSON-configured tax and compliance rules for cash, TDS, GST ITC, and large transaction checks."),
]


ANALYSIS_TABLES = [
    (
        "Pipeline Phase Matrix",
        ["Phase", "Source Component", "Primary Output"],
        [
            ("Data generation", "generate_dataset.py", "Synthetic ledger CSV files with Schedule III heads and GST fields"),
            ("Preprocessing", "preprocess.py and modules/preprocessing.py", "Cleaned data, TF-IDF features, amount scaling, encoded payment mode"),
            ("Training", "train_model.py", "Keras classifier, metrics, plots, anomaly model, serialized preprocessors"),
            ("Inference", "modules/ml_model.py", "Predicted category, confidence percentage, fallback-adjusted results"),
            ("Reasoning", "rule_engine.py and modules/financial_engine.py", "FinancialSummary, GST payable, P&L, Balance Sheet values"),
            ("Presentation", "app.py and modules/visualization.py", "Streamlit tabs, charts, dashboards, report downloads"),
        ],
        [1.35, 2.0, 3.15],
    ),
    (
        "Core Dependency Matrix",
        ["Library", "Used For", "Project Benefit"],
        [
            ("Pandas", "CSV/Excel ingestion and ledger transformations", "Fast table manipulation for accounting data"),
            ("NumPy", "Numerical operations and generated values", "Efficient arrays and synthetic value generation"),
            ("scikit-learn", "TF-IDF, scaling, label encoding, Isolation Forest", "Reliable ML preprocessing and anomaly detection"),
            ("TensorFlow/Keras", "Dense neural-network classifier", "Backpropagation-based category prediction"),
            ("Streamlit", "Application interface", "Quick interactive deployment for non-technical users"),
            ("Plotly", "Financial charts", "Interactive dashboards and business visualizations"),
            ("ReportLab", "PDF generation", "Client-ready CA report export"),
            ("Joblib", "Artifact serialization", "Reusable trained preprocessing and anomaly artifacts"),
        ],
        [1.4, 2.45, 2.65],
    ),
    (
        "Artifact Register",
        ["Artifact", "Produced By", "Purpose"],
        [
            ("ca_model.h5", "train_model.py", "Stores the trained transaction classifier"),
            ("tfidf_vectorizer.pkl", "preprocess.py", "Transforms descriptions into text features"),
            ("amount_scaler.pkl", "preprocess.py", "Normalizes transaction amount values"),
            ("payment_mode_encoder.pkl", "preprocess.py", "Encodes payment method labels"),
            ("anomaly_model.pkl", "train_model.py", "Supports Isolation Forest anomaly flagging"),
            ("metrics.json", "train_model.py", "Records accuracy, classes, input dimension, and classification report"),
            ("confusion_matrix.png", "train_model.py", "Visual evaluation of predicted versus actual classes"),
            ("training_history.png", "train_model.py", "Training and validation trend visualization"),
        ],
        [1.7, 1.6, 3.2],
    ),
    (
        "Input Dataset Fields",
        ["Field", "Meaning", "Used By"],
        [
            ("Date", "Transaction date", "Trend analysis and fiscal comparison"),
            ("Description", "Narrative transaction text", "TF-IDF and keyword fallback"),
            ("Amount", "Transaction value", "Financial summaries, scaling, anomaly detection"),
            ("Category", "Accounting class", "Training labels and rule-engine grouping"),
            ("Sub_Category", "Detailed account head", "Schedule III mapping and compliance rules"),
            ("GST_Percentage", "Tax slab", "GST breakdown and net GST payable"),
            ("GST_Amount", "Computed tax amount", "Input/output GST analysis"),
            ("Payment_Mode", "Cash, UPI, bank transfer, cheque, credit", "Compliance screening and feature engineering"),
        ],
        [1.55, 2.25, 2.7],
    ),
    (
        "Engine Responsibility Split",
        ["Layer", "Files", "Responsibility"],
        [
            ("UI orchestration", "app.py", "Handles upload, tabs, controls, and display state"),
            ("ML inference", "modules/ml_model.py", "Loads artifacts and predicts categories"),
            ("Accounting computation", "rule_engine.py, modules/financial_engine.py", "Builds summaries, ratios, cash-flow proxy, and validations"),
            ("Compliance", "modules/compliance_engine.py, rules/compliance_rules.json", "Generates statutory warnings and risk score"),
            ("Visualization", "modules/visualization.py", "Creates Plotly charts and interpretations"),
            ("Export", "modules/report_generator.py", "Builds downloadable professional PDF reports"),
        ],
        [1.4, 2.25, 2.85],
    ),
    (
        "Model Architecture Summary",
        ["Layer/Element", "Configuration", "Reason"],
        [
            ("Input", "504 features", "Combines TF-IDF text features with numeric/context features"),
            ("Dense block 1", "512 units with batch normalization and dropout", "Learns broad transaction patterns while controlling overfit"),
            ("Dense block 2", "256 units with batch normalization and dropout", "Refines category separation"),
            ("Dense blocks 3-4", "128 and 64 units", "Compresses learned representation before output"),
            ("Output", "Softmax over four classes", "Returns probabilities for Expense, Income, Asset, Liability"),
            ("Optimizer", "Adam", "Efficient gradient-based backpropagation training"),
            ("Loss", "Sparse categorical cross-entropy", "Appropriate for integer-encoded multi-class labels"),
        ],
        [1.55, 2.35, 2.6],
    ),
    (
        "Hybrid Prediction Control",
        ["Condition", "Action", "Reason"],
        [
            ("Artifacts missing", "Disable AI prediction", "Avoid misleading output when model files are unavailable"),
            ("Model confidence >= 0.70", "Use neural-network category", "Prediction is considered reliable enough for analysis"),
            ("Model confidence < 0.70", "Apply keyword fallback", "Accounting terms can rescue uncertain classifications"),
            ("Anomaly model available", "Add anomaly flag", "Identifies unusual records for review"),
            ("No category column", "Use Predicted_Category for downstream analysis", "Allows uploaded raw ledgers to be processed"),
        ],
        [1.8, 1.95, 2.75],
    ),
    (
        "Financial Computation Outputs",
        ["Output", "Source Logic", "Displayed In"],
        [
            ("Total Income", "Income-category accumulation", "Dashboard, report preview, PDF"),
            ("Total Expense", "Expense-category accumulation", "Dashboard, P&L, visualizations"),
            ("Net Profit/Loss", "Income minus expenses and provisions", "KPI cards and report"),
            ("GST Payable", "Output GST less input GST credit", "Dashboard, GST report section"),
            ("Total Assets", "Asset Schedule III heads", "Balance Sheet validation"),
            ("Total Liabilities", "Liability and equity Schedule III heads", "Balance Sheet validation"),
            ("Profit Margin", "Net result divided by income", "Health classification"),
        ],
        [1.65, 2.35, 2.5],
    ),
    (
        "Compliance Rule Register",
        ["Area", "Rule Basis", "Risk Signal"],
        [
            ("Cash expense", "Section 40A(3)", "Cash expense over Rs.10,000 may be disallowed"),
            ("Cash receipt/payment", "Section 269ST", "Cash transaction of Rs.2 lakh or more creates penalty risk"),
            ("Professional fees", "Section 194J", "TDS review needed above threshold"),
            ("Contractor payments", "Section 194C", "TDS review needed above threshold"),
            ("Blocked ITC", "Section 17(5) CGST", "Certain expenses cannot claim input credit"),
            ("High income", "Section 208", "Advance tax installment review may be required"),
        ],
        [1.45, 1.8, 3.25],
    ),
    (
        "Dashboard Tab Map",
        ["Tab", "Purpose", "Key Output"],
        [
            ("Data Preview", "Inspect uploaded ledger and quality", "Rows, columns, missing values, downloadable processed CSV"),
            ("AI Predictions", "Run or review classification", "Predicted category, confidence, anomaly flags, model performance"),
            ("Dashboard", "Summarize financial position", "KPIs, validations, income/expense breakdowns, cash flow"),
            ("Visualizations", "Explore transaction patterns", "Sankey, network, sunburst, monthly trend"),
            ("CA Insights", "Review compliance and advisory signals", "Score gauge, critical alerts, warnings, information"),
            ("Report", "Generate formal output", "14-section PDF and preview tables"),
        ],
        [1.45, 2.5, 2.55],
    ),
    (
        "Visualization Purpose Matrix",
        ["Chart", "Data Relationship", "CA Interpretation"],
        [
            ("Sankey flow", "Income sources to categories, sub-categories, GST/net", "Shows where money enters and leaves the business"),
            ("Network graph", "Category hubs, sub-categories, payment modes", "Highlights spending concentration and cash-channel risks"),
            ("Sunburst", "Category, sub-category, payment mode hierarchy", "Shows GST intensity and drill-down composition"),
            ("Monthly trend", "Income and expense over time", "Reveals seasonality, volatility, and margin pressure"),
            ("Compliance gauge", "Weighted risk score", "Condenses warnings into a reviewable score"),
        ],
        [1.45, 2.35, 2.7],
    ),
    (
        "Generated PDF Report Map",
        ["Section Group", "Examples", "Business Value"],
        [
            ("Financial statements", "P&L, Balance Sheet, GST analysis", "Presents accounting outputs in client-ready form"),
            ("Analysis", "Comparative analysis, cash flow, ratios, expenses", "Explains performance and operational pressure points"),
            ("Controls", "Compliance analysis, anomaly detection", "Directs review toward statutory and unusual-transaction risk"),
            ("Advisory", "CA recommendations, conclusion, disclaimer", "Turns data into action-oriented guidance"),
        ],
        [1.55, 2.55, 2.4],
    ),
    (
        "Testing Checklist",
        ["Test Area", "Check", "Expected Result"],
        [
            ("Pipeline", "Run dataset generation, preprocessing, and training", "Artifacts and metrics are produced"),
            ("Upload", "Load CSV and Excel files", "Dataframe appears without schema errors"),
            ("Prediction", "Run classification on test datasets", "Predicted_Category and confidence columns are created"),
            ("Rules", "Run dashboard computation", "FinancialSummary values and validations render"),
            ("Compliance", "Trigger cash/TDS/GST cases", "Appropriate alert severity appears"),
            ("Visualization", "Open each chart tab", "Charts render and remain interpretable"),
            ("Report", "Generate PDF", "Download button returns complete report bytes"),
        ],
        [1.35, 2.55, 2.6],
    ),
    (
        "Risk and Limitation Matrix",
        ["Limitation", "Impact", "Mitigation"],
        [
            ("Synthetic training data", "May not represent every real ledger pattern", "Retrain or fine-tune with validated real data"),
            ("Changing tax provisions", "Compliance rules can become outdated", "Maintain versioned rule JSON and periodic review"),
            ("Low-confidence text", "Ambiguous descriptions may be misclassified", "Expose confidence and fallback logic for manual review"),
            ("No authentication", "Local demo is not production-secure", "Add users, roles, audit logs, and encryption"),
            ("No source invoice verification", "Ledger-level checks cannot prove document authenticity", "Integrate invoice OCR and document matching"),
        ],
        [1.75, 2.25, 2.5],
    ),
    (
        "Future Scope Matrix",
        ["Enhancement", "Description", "Expected Benefit"],
        [
            ("Database backend", "Persist ledgers, reports, and audit history", "Supports multi-company usage"),
            ("OCR integration", "Read invoices and vouchers", "Improves source-document validation"),
            ("Rule versioning", "Track compliance rule changes by law and date", "Improves audit defensibility"),
            ("GST return reconciliation", "Compare ledger GST with returns", "Reduces indirect-tax mismatch risk"),
            ("Role-based access", "Separate client, accountant, and admin permissions", "Improves operational security"),
            ("Automated tests", "Unit and integration checks for modules", "Improves maintainability"),
        ],
        [1.6, 2.45, 2.45],
    ),
]


PAGE_TOPICS = [
    ("Executive Summary", "The project delivers an AI-assisted Chartered Accountant workflow for classifying transactions, computing Schedule III-style summaries, checking compliance exposure, visualizing financial behavior, and producing client-ready reports."),
    ("Problem Statement", "Manual ledger review is repetitive, error-prone, and slow when transaction volume grows. The project addresses classification, financial summarization, compliance screening, anomaly review, and report preparation in one workflow."),
    ("Project Objectives", "The objectives are to generate realistic financial data, preprocess it for machine learning, train a neural classifier, combine it with symbolic reasoning, and expose the results through a practical Streamlit interface."),
    ("Scope of the System", "The system covers transaction ingestion, category prediction, GST calculation, Profit and Loss computation, Balance Sheet reconstruction, compliance scoring, visualization, and report generation."),
    ("Technology Stack", "The implementation uses Python, Pandas, NumPy, scikit-learn, TensorFlow/Keras, Streamlit, Plotly, Joblib, ReportLab, Matplotlib, Seaborn, and JSON-based rules."),
    ("Repository Organization", "The repository separates root-level training scripts from production modules. The modules package holds reusable runtime logic, while app.py acts as the UI layer."),
    ("Configuration Strategy", "config.py centralizes directories, artifact paths, category labels, thresholds, model parameters, palette values, and application metadata so modules avoid scattered constants."),
    ("Data Generation Design", "generate_dataset.py creates synthetic ledgers with financial categories, sub-categories, Schedule III heads, GST percentages, payment modes, vendor names, dates, and anomaly flags."),
    ("Company Profiles", "The generator supports business profiles such as generic, trading, manufacturing, service, and startup. Each profile changes category distributions and makes test data more realistic."),
    ("Transaction Schema", "Transactions include date, description, amount, transaction type, category, sub-category, Schedule III head, GST percentage, GST amount, payment mode, vendor/client, and anomaly flag."),
    ("GST Modeling", "GST is modeled through slab rates and exemptions. Income and expense records can carry GST, while several asset, liability, equity, reserve, and tax-provision heads are exempt."),
    ("Noise Injection", "The dataset intentionally includes missing values and spelling variants. This makes preprocessing and model inference closer to a real accounting data-cleaning problem."),
    ("Training Dataset Summary", "The cleaned training dataset contains hundreds of thousands of records with four target classes: Expense, Income, Asset, and Liability."),
    ("Test Dataset Profiles", "Five 12,000-row test datasets represent generic, manufacturing, service, startup, and trading businesses. This supports scenario-based evaluation and demonstration."),
    ("Preprocessing Pipeline", "preprocess.py normalizes text, encodes payment modes, scales amount, creates calendar features, transforms descriptions with TF-IDF, and serializes the feature matrix."),
    ("Text Normalization", "Transaction descriptions are lowercased, punctuation is removed, and whitespace is normalized. Missing or invalid text becomes a safe unknown token."),
    ("TF-IDF Feature Engineering", "The model uses 500 TF-IDF features with unigram and bigram terms. This captures words such as salary, purchase, sales, loan, GST, rent, and invoice context."),
    ("Numerical Feature Engineering", "Amount is scaled with StandardScaler, payment mode is encoded, and date-derived month/weekend signals help inference capture context beyond text."),
    ("Target Encoding", "The four accounting categories map to numeric labels: Expense, Income, Asset, and Liability. The mapping is centralized to keep training and inference consistent."),
    ("Model Architecture", "train_model.py defines a dense feed-forward neural network with batch normalization, dropout, and softmax output for four-class transaction classification."),
    ("Backpropagation Training", "The classifier is trained with Adam and sparse categorical cross-entropy. Early stopping and validation metrics reduce the risk of overfitting."),
    ("Model Artifacts", "Training writes ca_model.h5, TF-IDF vectorizer, amount scaler, payment mode encoder, feature data, metrics JSON, confusion matrix, training history, and anomaly model."),
    ("Model Performance", "The metrics artifact reports 99.72 percent test accuracy, 504 input features, and strong precision/recall across all four categories."),
    ("Classification Report", "Expense, Income, Asset, and Liability each show F1 scores near 0.997, indicating that the synthetic data patterns are learned reliably by the neural network."),
    ("Hybrid Inference", "modules/ml_model.py loads artifacts, builds runtime feature matrices, predicts categories, and falls back to keyword rules when model confidence is below 0.70."),
    ("Keyword Fallback Rules", "Low-confidence predictions can be corrected by lightweight keyword matching for salary, rent, sales, purchase, loan, GST, asset, equity, and similar accounting language."),
    ("Anomaly Modeling", "The training flow also persists an Isolation Forest model, while modules/anomaly_detection.py adds statistical analysis using amount distributions and summary counts."),
    ("Backward Chaining Concept", "The rule engine starts with top-level goals such as Profit and Loss, GST payable, and Balance Sheet reconciliation, then proves each goal from transaction facts."),
    ("Financial Facts", "Each row becomes a fact with category, sub-category, Schedule III head, amount, GST, payment mode, and descriptive metadata. These facts are the knowledge base."),
    ("Profit and Loss Goals", "The engine proves revenue from operations, other income, material cost, employee benefits, finance costs, depreciation, other expenses, tax provision, and net result."),
    ("Balance Sheet Goals", "The engine proves shareholders funds, non-current liabilities, current liabilities, non-current assets, current assets, total liabilities, and total assets."),
    ("GST Goals", "GST collected on income and GST paid on eligible expenses are accumulated to calculate net GST payable. This is used in both dashboards and reports."),
    ("Reconciliation Goal", "The balance-sheet validation checks whether total assets equal equity and liabilities, then reports the gap and a practical explanatory note."),
    ("Financial Health Rules", "Profit margin thresholds classify health as Strong, Moderate, Weak, or Loss-Making. These labels drive UI badges, validations, insights, and report text."),
    ("Cash Flow Proxy", "The financial engine estimates operating cash flow using net profit, depreciation, finance costs, working-capital proxy, and available cash signals."),
    ("Comparative Analysis", "When dates span multiple fiscal periods, the system computes year-over-year revenue, expense, profit, and trend metrics."),
    ("Compliance Rule Design", "Compliance rules are stored in JSON with rule name, condition, severity, section, and message. This keeps regulatory checks editable without modifying Python code."),
    ("Cash Payment Compliance", "Rules flag Section 40A(3) cash expenses over Rs. 10,000 and Section 269ST cash transactions of Rs. 2 lakh or more."),
    ("TDS Compliance", "Rules flag professional fees under Section 194J and contractor payments under Section 194C when thresholds are crossed."),
    ("GST ITC Compliance", "The engine flags blocked input tax credit categories such as staff welfare, food and beverages, vehicle repair, and health insurance."),
    ("Large Transaction Review", "Very large transactions are highlighted for documentation, authorization, and governance review."),
    ("Compliance Scoring", "Critical and warning insights are weighted into a compliance score. The score starts at 100 and deducts risk points based on issue severity."),
    ("Data Quality Validation", "The validation layer checks missing values, duplicate rows, negative or zero amounts, missing dates, invalid categories, and other ingestion risks."),
    ("Utility Layer", "utils/helpers.py provides Indian currency formatting, PDF-safe currency formatting, emoji stripping, text cleaning, data cleaning, and data quality scoring."),
    ("Logging Strategy", "utils/logger.py configures console and rotating-file logging under logs/ca_suite.log so long-running pipeline and app behavior can be reviewed."),
    ("Streamlit Application", "app.py is intentionally described as a pure UI layer. It imports business logic from modules and orchestrates upload, prediction, dashboard, insight, and report views."),
    ("Sidebar Workflow", "The sidebar handles theme toggle, file upload, company details, financial year, AI prediction option, and model status with test accuracy."),
    ("Data Preview Tab", "The first tab previews the uploaded dataset, displays key metrics, allows downloading the processed CSV, and renders the data quality panel."),
    ("AI Predictions Tab", "The prediction tab can run classification, display prediction counts, confidence values, anomalies, model accuracy, confusion matrix, and training history."),
    ("Dashboard Tab", "The dashboard tab runs the backward chaining engine and presents health status, KPI cards, validations, breakdowns, cash-flow insight, and comparative analysis."),
    ("Visualization Tab", "The visualization tab presents Sankey flow, category/payment network graph, GST-intensity sunburst hierarchy, and monthly income-versus-expense trend."),
    ("CA Insights Tab", "The insights tab combines compliance score gauge, critical alerts, warnings, and informational recommendations generated from financial and compliance engines."),
    ("Report Tab", "The report tab generates a 14-section professional CA financial report as a downloadable PDF using modules/report_generator.py."),
    ("PDF Report Sections", "The generated PDF contains executive summary, Profit and Loss, Balance Sheet, comparative analysis, cash flow, GST, ratios, expenses, charts, compliance, anomalies, recommendations, conclusion, and disclaimer."),
    ("Visualization Engineering", "Plotly figures use the configured palette and contextual interpretations. The charts are designed to explain the accounting implication, not only show numbers."),
    ("Security and Safety Notes", "The system is a decision-support tool. It should not replace professional judgment, statutory audit procedures, source-document verification, or final CA sign-off."),
    ("Testing Approach", "The project can be tested by running the pipeline, checking artifact generation, loading all sample datasets, validating model metrics, and exercising each Streamlit tab."),
    ("Known Limitations", "The dataset is synthetic, legal thresholds may change, keyword fallback is intentionally lightweight, and report conclusions depend on the quality of uploaded ledger data."),
    ("Future Enhancements", "Future work could add database storage, real invoice ingestion, OCR, authentication, audit trails, rule versioning, GST return reconciliation, and automated unit tests."),
    ("Deployment Considerations", "The app can run locally with Streamlit, while model artifacts and datasets should be versioned carefully. Production deployment would require privacy and access controls."),
    ("Conclusion", "CA Intelligence Suite demonstrates a hybrid AI approach that combines neural-network classification, symbolic backward chaining, compliance rules, visual analytics, and report automation."),
]


def add_cover(doc):
    para(doc, "PROJECT REPORT", before=40, after=8, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=MUTED)
    para(doc, "CA Intelligence Suite", after=8, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26, color=NAVY)
    para(doc, "AI-Powered Chartered Accountant Assistant", after=24, align=WD_ALIGN_PARAGRAPH.CENTER, size=15, color=MUTED)
    rows = [
        ("Project Type", "Machine Learning + Symbolic AI + Streamlit Application"),
        ("Prepared From", "Local source code, datasets, configuration, rules, and model artifacts"),
        ("Prepared On", date.today().strftime("%B %d, %Y")),
        ("Repository", str(ROOT)),
    ]
    add_table(doc, ["Field", "Value"], rows, widths=[1.6, 4.9])
    para(doc, "This document is a code-grounded academic project report. It explains the design, implementation, workflow, algorithms, evaluation, limitations, and future scope of the CA Intelligence Suite project.", after=10, line=1.25)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    for i, (title, _) in enumerate(PAGE_TOPICS, start=1):
        para(doc, f"{i}. {title}", after=2)
    doc.add_page_break()


def add_project_tables(doc, metrics):
    add_heading(doc, "Source Code Inventory", 1)
    add_table(doc, ["File", "Role in Project"], PROJECT_FILES, widths=[1.9, 4.6])
    doc.add_page_break()

    add_heading(doc, "Dataset Inventory", 1)
    rows = dataset_overview()
    add_table(doc, ["Dataset", "Rows", "Columns", "Category Distribution"], rows, widths=[2.0, 0.8, 0.8, 2.9])
    para(doc, "The training dataset is much larger than the demonstration datasets and includes engineered columns such as Description_clean, Month, Is_Weekend, encoded payment mode, encoded category, Log_Amount, and Amount_scaled.", line=1.25)
    doc.add_page_break()

    add_heading(doc, "Model Metrics", 1)
    rows = [
        ("Test accuracy", f"{metrics.get('test_accuracy', 0) * 100:.2f}%"),
        ("Epochs run", str(metrics.get("epochs_run", "Not available"))),
        ("Input dimension", str(metrics.get("input_dim", "Not available"))),
        ("Classes", ", ".join(metrics.get("classes", []))),
    ]
    add_table(doc, ["Metric", "Value"], rows, widths=[2.2, 4.3])
    report = metrics.get("classification_report", {})
    class_rows = []
    for cls in ["Expense", "Income", "Asset", "Liability"]:
        rec = report.get(cls, {})
        class_rows.append([
            cls,
            f"{rec.get('precision', 0):.4f}",
            f"{rec.get('recall', 0):.4f}",
            f"{rec.get('f1-score', 0):.4f}",
            f"{rec.get('support', 0):,.0f}",
        ])
    add_table(doc, ["Class", "Precision", "Recall", "F1 Score", "Support"], class_rows, widths=[1.3, 1.2, 1.2, 1.2, 1.6])
    doc.add_page_break()

    add_heading(doc, "Detailed Design Tables", 1)
    for title, headers, rows, widths in ANALYSIS_TABLES:
        add_heading(doc, title, 2)
        add_table(doc, headers, rows, widths=widths)
    doc.add_page_break()


def topic_details(title):
    base = [
        f"In the source code, this topic is implemented as part of a layered architecture where the Streamlit interface delegates domain behavior to reusable modules. The design keeps project explanation, prediction, accounting computation, compliance screening, visualization, and report generation separated enough to be understood and tested independently.",
        f"For the project report, {title.lower()} matters because it connects the accounting domain with the software implementation. The code does not only classify rows; it converts transaction data into decisions a CA user can inspect, challenge, and export.",
        "The implementation favors practical clarity: data enters as CSV or Excel, preprocessing standardizes it, machine learning predicts categories when needed, the rule engine computes financial summaries, and the UI presents the outcome through metrics, validations, charts, insights, and downloadable reports.",
    ]
    return base


def add_topic_page(doc, num, title, summary):
    add_heading(doc, f"{num}. {title}", 1)
    para(doc, summary, bold=True, color=NAVY, line=1.25, after=10)
    for paragraph in topic_details(title):
        para(doc, paragraph, line=1.25, after=7)
    add_heading(doc, "Implementation Notes", 2)
    add_bullets(doc, [
        "Relevant behavior is visible in the project files and is reflected in the report narrative.",
        "The implementation combines data engineering, machine learning, accounting rules, and UI orchestration.",
        "The design is suitable for demonstration, academic evaluation, and extension into a more production-ready audit assistant.",
    ])
    add_heading(doc, "Inputs", 3)
    add_bullets(doc, [
        "Transaction records, configuration values, serialized artifacts, or computed summaries depending on the module being discussed.",
        "The application keeps user-facing upload data separate from reusable project constants and generated model files.",
    ])
    add_heading(doc, "Outputs", 3)
    add_bullets(doc, [
        "Cleaned datasets, predictions, accounting totals, compliance insights, charts, logs, or generated report artifacts.",
        "Outputs are designed to be readable by both technical reviewers and accounting-domain users.",
    ])
    add_heading(doc, "Project Relevance", 2)
    para(doc, "This section supports the overall project by explaining why the feature exists, how it is connected to the rest of the pipeline, and what value it provides to a Chartered Accountant or finance reviewer.", line=1.25)
    doc.add_page_break()


def add_appendices(doc):
    add_heading(doc, "Appendix A: Compliance Rules", 1)
    rules_path = ROOT / "rules" / "compliance_rules.json"
    rules = json.loads(rules_path.read_text(encoding="utf-8")) if rules_path.exists() else []
    rows = [[r["section"], r["severity"], r["condition"], r["message"]] for r in rules]
    add_table(doc, ["Section", "Severity", "Condition", "Message"], rows, widths=[0.9, 0.9, 2.0, 2.7])
    doc.add_page_break()

    add_heading(doc, "Appendix B: Suggested Viva Questions", 1)
    add_numbered(doc, [
        "Why did the project combine neural-network classification with a symbolic backward chaining rule engine?",
        "How does TF-IDF help in transaction classification?",
        "What are the four output classes and why are they important for accounting analysis?",
        "How is GST payable computed from income and expense transactions?",
        "What is the role of the compliance_rules.json file?",
        "Why is synthetic data useful for this project?",
        "What are the main limitations of using synthetic data?",
        "How does the system handle low-confidence ML predictions?",
        "Which parts of the application should be tested before deployment?",
        "How could the project be extended for real-world audit usage?",
    ])
    doc.add_page_break()

    add_heading(doc, "Appendix C: Glossary", 1)
    add_table(doc, ["Term", "Meaning"], [
        ("Backward chaining", "A reasoning technique that begins with a goal and works backward to prove it from available facts."),
        ("TF-IDF", "A text feature method that scores words by importance within documents and the corpus."),
        ("GST", "Goods and Services Tax, represented in the project through slab percentages and input/output GST amounts."),
        ("Schedule III", "Companies Act 2013 format used as the accounting structure for financial statements."),
        ("Isolation Forest", "An unsupervised anomaly detection algorithm used to flag unusual transactions."),
        ("Compliance score", "A risk-weighted score derived from critical, warning, and informational insights."),
    ], widths=[1.8, 4.7])
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    metrics = load_metrics()
    add_project_tables(doc, metrics)
    for i, (title, summary) in enumerate(PAGE_TOPICS, start=1):
        add_topic_page(doc, i, title, summary)
    add_appendices(doc)
    REPORT_PATH.parent.mkdir(exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
